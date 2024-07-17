import streamlit as st
import numpy as np
import pickle
import os
from dotenv import load_dotenv
from docx import Document
import google.generativeai as genai
from io import BytesIO

# Load environment variables from .env file
load_dotenv()

# List of diseases and symptoms
diseases = [
    '(vertigo) Paroymsal Positional Vertigo', 'AIDS', 'Acne', 'Alcoholic hepatitis', 'Allergy',
    'Arthritis', 'Bronchial Asthma', 'Cervical spondylosis', 'Chicken pox', 'Chronic cholestasis',
    'Common Cold', 'Dengue', 'Diabetes', 'Dimorphic hemmorhoids(piles)', 'Drug Reaction',
    'Fungal infection', 'GERD', 'Gastroenteritis', 'Heart attack', 'Hepatitis B', 'Hepatitis C',
    'Hepatitis D', 'Hepatitis E', 'Hypertension', 'Hyperthyroidism', 'Hypoglycemia', 'Hypothyroidism',
    'Impetigo', 'Jaundice', 'Malaria', 'Migraine', 'Osteoarthristis', 'Paralysis (brain hemorrhage)',
    'Peptic ulcer diseae', 'Pneumonia', 'Psoriasis', 'Tuberculosis', 'Typhoid', 'Urinary tract infection',
    'Varicose veins', 'hepatitis A'
]

symptoms = ['itching', 'skin_rash', 'nodal_skin_eruptions', 'continuous_sneezing', 'shivering', 'chills', 'joint_pain', 'stomach_pain', 'acidity', 'ulcers_on_tongue', 'muscle_wasting', 'vomiting', 'burning_micturition', 'spotting_urination', 'fatigue', 'weight_gain', 'anxiety', 'cold_hands_and_feets', 'mood_swings', 'weight_loss', 'restlessness', 'lethargy', 'patches_in_throat', 'irregular_sugar_level', 'cough', 'high_fever', 'sunken_eyes', 'breathlessness', 'sweating', 'dehydration', 'indigestion', 'headache', 'yellowish_skin', 'dark_urine', 'nausea', 'loss_of_appetite', 'pain_behind_the_eyes', 'back_pain', 'constipation', 'abdominal_pain', 'diarrhoea', 'mild_fever', 'yellow_urine', 'yellowing_of_eyes', 'acute_liver_failure', 'fluid_overload', 'swelling_of_stomach', 'swelled_lymph_nodes', 'malaise', 'blurred_and_distorted_vision', 'phlegm', 'throat_irritation', 'redness_of_eyes', 'sinus_pressure', 'runny_nose', 'congestion', 'chest_pain', 'weakness_in_limbs', 'fast_heart_rate', 'pain_during_bowel_movements', 'pain_in_anal_region', 'bloody_stool', 'irritation_in_anus', 'neck_pain', 'dizziness', 'cramps', 'bruising', 'obesity', 'swollen_legs', 'swollen_blood_vessels', 'puffy_face_and_eyes', 'enlarged_thyroid', 'brittle_nails', 'swollen_extremeties', 'excessive_hunger', 'extra_marital_contacts', 'drying_and_tingling_lips', 'slurred_speech', 'knee_pain', 'hip_joint_pain', 'muscle_weakness', 'stiff_neck', 'swelling_joints', 'movement_stiffness', 'spinning_movements', 'loss_of_balance', 'unsteadiness', 'weakness_of_one_body_side', 'loss_of_smell', 'bladder_discomfort', 'foul_smell_ofurine', 'continuous_feel_of_urine', 'passage_of_gases', 'internal_itching', 'toxic_look_(typhos)', 'depression', 'irritability', 'muscle_pain', 'altered_sensorium', 'red_spots_over_body', 'belly_pain', 'abnormal_menstruation', 'dischromic_patches', 'watering_from_eyes', 'increased_appetite', 'polyuria', 'family_history', 'mucoid_sputum', 'rusty_sputum', 'lack_of_concentration', 'visual_disturbances', 'receiving_blood_transfusion', 'receiving_unsterile_injections', 'coma', 'stomach_bleeding', 'distention_of_abdomen', 'history_of_alcohol_consumption', 'blood_in_sputum', 'prominent_veins_on_calf', 'palpitations', 'painful_walking', 'pus_filled_pimples', 'blackheads', 'scurring', 'skin_peeling', 'silver_like_dusting', 'small_dents_in_nails', 'inflammatory_nails', 'blister', 'red_sore_around_nose', 'yellow_crust_ooze', 'prognosis', ' skin_rash', ' nodal_skin_eruptions', ' dischromic _patches', ' continuous_sneezing', ' shivering', ' chills', ' watering_from_eyes', ' stomach_pain', ' acidity', ' ulcers_on_tongue', ' vomiting', ' cough', ' chest_pain', ' yellowish_skin', ' nausea', ' loss_of_appetite', ' abdominal_pain', ' yellowing_of_eyes', ' burning_micturition', ' spotting_ urination', ' passage_of_gases', ' internal_itching', ' indigestion', ' muscle_wasting', ' patches_in_throat', ' high_fever', ' extra_marital_contacts', ' fatigue', ' weight_loss', ' restlessness', ' lethargy', ' irregular_sugar_level', ' blurred_and_distorted_vision', ' obesity', ' excessive_hunger', ' increased_appetite', ' polyuria', ' sunken_eyes', ' dehydration', ' diarrhoea', ' breathlessness', ' family_history', ' mucoid_sputum', ' headache', ' dizziness', ' loss_of_balance', ' lack_of_concentration', ' stiff_neck', ' depression', ' irritability', ' visual_disturbances', ' back_pain', ' weakness_in_limbs', ' neck_pain', ' weakness_of_one_body_side', ' altered_sensorium', ' dark_urine', ' sweating', ' muscle_pain', ' mild_fever', ' swelled_lymph_nodes', ' malaise', ' red_spots_over_body', ' joint_pain', ' pain_behind_the_eyes', ' constipation', ' toxic_look_(typhos)', ' belly_pain', ' yellow_urine', ' receiving_blood_transfusion', ' receiving_unsterile_injections', ' coma', ' stomach_bleeding', ' acute_liver_failure', ' swelling_of_stomach', ' distention_of_abdomen', ' history_of_alcohol_consumption', ' fluid_overload', ' phlegm', ' blood_in_sputum', ' throat_irritation', ' redness_of_eyes', ' sinus_pressure', ' runny_nose', ' congestion', ' loss_of_smell', ' fast_heart_rate', ' rusty_sputum', ' pain_during_bowel_movements', ' pain_in_anal_region', ' bloody_stool', ' irritation_in_anus', ' cramps', ' bruising', ' swollen_legs', ' swollen_blood_vessels', ' prominent_veins_on_calf', ' weight_gain', ' cold_hands_and_feets', ' mood_swings', ' puffy_face_and_eyes', ' enlarged_thyroid', ' brittle_nails', ' swollen_extremeties', ' abnormal_menstruation', ' muscle_weakness', ' anxiety', ' slurred_speech', ' palpitations', ' drying_and_tingling_lips', ' knee_pain', ' hip_joint_pain', ' swelling_joints', ' painful_walking', ' movement_stiffness', ' spinning_movements', ' unsteadiness', ' pus_filled_pimples', ' blackheads', ' scurring', ' bladder_discomfort', ' foul_smell_of urine', ' continuous_feel_of_urine', ' skin_peeling', ' silver_like_dusting', ' small_dents_in_nails', ' inflammatory_nails', ' blister', ' red_sore_around_nose', ' yellow_crust_ooze']

# Load the XGBoost model
@st.cache(allow_output_mutation=True)
def load_model():
    try:
        with open('xgb.pkl', 'rb') as model_file:
            model = pickle.load(model_file)
        st.success("Model loaded successfully")
        return model
    except Exception as e:
        st.error(f"Error loading model: {str(e)}")
        raise

# Initialize Gemini model
def initialize_gemini():
    try:
        api_key = os.getenv("GEMINI_API_KEY")
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-1.5-pro')
        return model
    except Exception as e:
        st.error(f"Error initializing Gemini model: {str(e)}")
        raise

# Function to predict top 5 diseases using XGBoost model
def predict_disease(symptoms_list, model):
    try:
        input_data = np.zeros(len(symptoms), dtype=int)
        for symptom in symptoms_list:
            if symptom in symptoms:
                input_data[symptoms.index(symptom)] = 1
        input_data = input_data.reshape(1, -1)

        probs = model.predict_proba(input_data)[0]
        top5_prob_indices = np.argsort(probs)[::-1][:5]
        top5_predicted_diseases = [(diseases[i], probs[i]) for i in top5_prob_indices]

        return top5_predicted_diseases
    except Exception as e:
        st.error(f"Error predicting disease: {str(e)}")
        raise

# Generate medical advice using Gemini model
def generate_advice(top5_predicted_diseases, symptoms_list, model):
    try:
        advice = "Generating advice...\n"
        for disease, probability in top5_predicted_diseases:
            prompt = f"Provide advice for {disease} based on symptoms: {', '.join(symptoms_list)}"
            response = model.generate_content(prompt)
            if response.text:
                advice += response.text.strip() + "\n\n"
            else:
                advice += f"Unable to generate advice for {disease}\n\n"
        return advice.strip()
    except Exception as e:
        st.error(f"Error generating advice: {str(e)}")
        raise

# Chatbot interaction function
def chatbot_interaction(model, user_input):
    try:
        response = model.generate_content(user_input)
        return response.text.strip()
    except Exception as e:
        st.error(f"Error in chatbot interaction: {str(e)}")
        return f"I apologize, but I'm having trouble processing your request at the moment. Your question was: '{user_input}'. Please try rephrasing your question or ask something else."

# Generate full report
def generate_full_report(top5_predicted_diseases, symptoms_list, advice):
    doc = Document()
    doc.add_heading("Disease Prediction and Medical Advice Report", 0)
    
    doc.add_heading("Symptoms", level=1)
    doc.add_paragraph(", ".join(symptoms_list))
    
    doc.add_heading("Top 5 Predicted Diseases", level=1)
    for disease, probability in top5_predicted_diseases:
        doc.add_paragraph(f"{disease} (Probability: {probability:.2f})")
    
    doc.add_heading("Medical Advice", level=1)
    doc.add_paragraph(advice)
    
    # Save the document to a BytesIO object
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io

# Main function
def main():
    st.title("HealthPredict AI-Powered Disease Forecasting and Medical Aid System")

    # Load XGBoost model
    model = load_model()

    # Initialize Gemini model
    gemini_model = initialize_gemini()

    # Create tabs for Disease Prediction and Chatbot
    tab1, tab2 = st.tabs(["Disease Prediction", "Medical Chatbot"])

    with tab1:
        st.markdown("Enter symptoms separated by commas.")
        symptoms_input = st.text_input("Enter Symptoms")

        if st.button("Predict"):
            symptoms_list = [symptom.strip() for symptom in symptoms_input.split(',')]

            top5_predicted_diseases = predict_disease(symptoms_list, model)

            st.subheader("Top 5 Predicted Diseases:")
            for disease, probability in top5_predicted_diseases:
                st.write(f"{disease} (Probability: {probability:.2f})")

            st.subheader("Medical Advice:")
            st.write("Generating advice...")

            advice = generate_advice(top5_predicted_diseases, symptoms_list, gemini_model)
            st.write(advice)

            # Generate full report
            report = generate_full_report(top5_predicted_diseases, symptoms_list, advice)
            
            # Provide download button
            st.download_button(
                label="Download Full Report",
                data=report,
                file_name="medical_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    with tab2:
        st.subheader("Medical Chatbot")
        st.write("Ask any medical questions, and our AI assistant will provide information.")

        if "messages" not in st.session_state:
            st.session_state.messages = []

        for message in st.session_state.messages:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])

        if prompt := st.chat_input("What is your question?"):
            st.chat_message("user").markdown(prompt)
            st.session_state.messages.append({"role": "user", "content": prompt})

            response = chatbot_interaction(gemini_model, prompt)

            with st.chat_message("assistant"):
                st.markdown(response)
            st.session_state.messages.append({"role": "assistant", "content": response})

if __name__ == "__main__":
    main()